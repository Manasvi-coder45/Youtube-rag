import os
import io
import re
import requests
import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain_core.runnables import RunnableParallel, RunnablePassthrough, RunnableLambda
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv

load_dotenv()

# ══════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════
st.set_page_config(
    page_title="TubeIQ — Video Intelligence",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════
# GLOBAL CSS  — Warm Editorial Light Theme
# Palette: cream bg · deep teal accent · amber highlight
# Fonts:   Playfair Display (headings) + DM Sans (body)
# ══════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;800;900&family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap');

/* ── Tokens ── */
:root {
  --cream:    #faf7f2;
  --cream2:   #f3ede3;
  --cream3:   #ede5d8;
  --teal:     #1a6b6b;
  --teal2:    #2a8f8f;
  --teal-lt:  #e8f4f4;
  --amber:    #d97706;
  --amber-lt: #fef3c7;
  --rose:     #be123c;
  --rose-lt:  #fff1f2;
  --ink:      #1c1917;
  --ink2:     #44403c;
  --ink3:     #78716c;
  --border:   #d6cfc4;
  --border2:  #c5bdb0;
  --shadow:   0 2px 16px rgba(28,25,23,.08);
  --shadow2:  0 4px 32px rgba(28,25,23,.12);
  --font-h:   'Playfair Display', Georgia, serif;
  --font-b:   'DM Sans', sans-serif;
  --font-m:   'DM Mono', monospace;
  --r:        12px;
}

/* ── Base ── */
html, body, [class*="css"] {
  font-family: var(--font-b);
  background: var(--cream);
  color: var(--ink);
}
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 0 2.2rem 5rem; max-width: 1360px; margin: 0 auto; }

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 5px; height: 5px; }
::-webkit-scrollbar-track { background: var(--cream2); }
::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 10px; }

/* ══════════════════════════════
   SIDEBAR
══════════════════════════════ */
[data-testid="stSidebar"] {
  background: var(--ink) !important;
  border-right: none !important;
}
[data-testid="stSidebar"] > div { padding: 1.8rem 1.4rem; }

/* sidebar text overrides */
[data-testid="stSidebar"] * { color: #d6d3d1 !important; }
[data-testid="stSidebar"] strong { color: #faf7f2 !important; }

[data-testid="stSidebar"] [data-testid="stTextInput"] input {
  background: #292524 !important;
  border: 1px solid #44403c !important;
  border-radius: 8px !important;
  color: #faf7f2 !important;
  font-family: var(--font-b) !important;
  font-size: .9rem !important;
  padding: .65rem 1rem !important;
}
[data-testid="stSidebar"] [data-testid="stTextInput"] input:focus {
  border-color: #2a8f8f !important;
  box-shadow: 0 0 0 3px rgba(42,143,143,.2) !important;
}
[data-testid="stSidebar"] [data-testid="stTextInput"] label {
  color: #a8a29e !important;
  font-size: .78rem !important;
}

/* sidebar buttons */
[data-testid="stSidebar"] .stButton > button {
  background: #292524 !important;
  color: #d6d3d1 !important;
  border: 1px solid #44403c !important;
  border-radius: 8px !important;
  font-family: var(--font-b) !important;
  font-size: .85rem !important;
  padding: .55rem 1rem !important;
  width: 100%;
  transition: all .18s;
}
[data-testid="stSidebar"] .stButton > button:hover {
  background: #3c3733 !important;
  border-color: #2a8f8f !important;
  color: #e7e5e4 !important;
  transform: translateY(-1px) !important;
}

/* sidebar form submit */
[data-testid="stSidebar"] [data-testid="stFormSubmitButton"] > button {
  background: var(--teal) !important;
  color: #fff !important;
  border: none !important;
  border-radius: 8px !important;
  font-family: var(--font-b) !important;
  font-weight: 600 !important;
  font-size: .88rem !important;
  padding: .65rem 1rem !important;
  width: 100%;
  transition: opacity .18s, transform .1s;
  letter-spacing: .01em;
}
[data-testid="stSidebar"] [data-testid="stFormSubmitButton"] > button:hover {
  opacity: .88 !important;
  transform: translateY(-1px) !important;
}

/* ══════════════════════════════
   MAIN AREA INPUTS
══════════════════════════════ */
[data-testid="stTextInput"] input {
  background: #fff !important;
  border: 1.5px solid var(--border) !important;
  border-radius: 10px !important;
  color: var(--ink) !important;
  font-family: var(--font-b) !important;
  font-size: .92rem !important;
  padding: .72rem 1rem !important;
  transition: border-color .2s, box-shadow .2s;
  box-shadow: 0 1px 4px rgba(28,25,23,.05);
}
[data-testid="stTextInput"] input:focus {
  border-color: var(--teal) !important;
  box-shadow: 0 0 0 3px rgba(26,107,107,.12) !important;
  outline: none !important;
}
[data-testid="stTextInput"] label {
  color: var(--ink3) !important;
  font-size: .8rem !important;
  font-weight: 500 !important;
}

/* ══════════════════════════════
   MAIN BUTTONS
══════════════════════════════ */
.stButton > button {
  background: #fff !important;
  color: var(--ink2) !important;
  border: 1.5px solid var(--border) !important;
  border-radius: 8px !important;
  font-family: var(--font-b) !important;
  font-weight: 500 !important;
  font-size: .85rem !important;
  padding: .55rem 1rem !important;
  transition: all .18s ease !important;
  width: 100%;
  box-shadow: 0 1px 3px rgba(28,25,23,.06);
}
.stButton > button:hover {
  background: var(--cream2) !important;
  border-color: var(--teal) !important;
  color: var(--teal) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 3px 8px rgba(28,25,23,.1) !important;
}

/* Main form submit — teal CTA */
[data-testid="stFormSubmitButton"] > button {
  background: var(--teal) !important;
  color: #fff !important;
  border: none !important;
  border-radius: 8px !important;
  font-family: var(--font-b) !important;
  font-weight: 600 !important;
  font-size: .88rem !important;
  padding: .65rem 1rem !important;
  width: 100%;
  box-shadow: 0 2px 8px rgba(26,107,107,.25);
  transition: opacity .18s, transform .1s;
}
[data-testid="stFormSubmitButton"] > button:hover {
  opacity: .88 !important;
  transform: translateY(-1px) !important;
}

/* ══════════════════════════════
   TABS
══════════════════════════════ */
[data-testid="stTabs"] {
  border-bottom: 2px solid var(--cream3);
  margin-bottom: 0;
}
[data-testid="stTabs"] button {
  font-family: var(--font-b) !important;
  font-size: .88rem !important;
  font-weight: 500 !important;
  color: var(--ink3) !important;
  padding: .65rem 1.3rem !important;
  border-radius: 0 !important;
  border: none !important;
  background: transparent !important;
  transition: color .18s;
}
[data-testid="stTabs"] button:hover { color: var(--teal) !important; }
[data-testid="stTabs"] button[aria-selected="true"] {
  color: var(--teal) !important;
  font-weight: 600 !important;
  border-bottom: 2.5px solid var(--teal) !important;
}
[data-testid="stTabs"] [data-testid="stTabsContent"] { padding-top: 1.6rem; }

/* ══════════════════════════════
   HERO
══════════════════════════════ */
.hero-wrap {
  padding: 3.5rem 0 2rem;
  text-align: center;
}
.hero-eyebrow {
  font-family: var(--font-b);
  font-size: .7rem;
  font-weight: 600;
  letter-spacing: .2em;
  text-transform: uppercase;
  color: var(--teal);
  margin-bottom: 1rem;
  opacity: .85;
}
.hero-title {
  font-family: var(--font-h);
  font-size: clamp(3.2rem, 6vw, 5.5rem);
  font-weight: 900;
  line-height: .95;
  letter-spacing: -.02em;
  color: var(--ink);
  margin-bottom: .8rem;
}
.hero-title span {
  color: var(--teal);
}
.hero-sub {
  font-size: 1.05rem;
  font-weight: 300;
  color: var(--ink3);
  max-width: 460px;
  margin: 0 auto;
  line-height: 1.7;
}

/* ══════════════════════════════
   CARDS & SURFACES
══════════════════════════════ */
.card {
  background: #fff;
  border: 1.5px solid var(--border);
  border-radius: var(--r);
  padding: 1.4rem;
  box-shadow: var(--shadow);
}
.feat-card {
  background: #fff;
  border: 1.5px solid var(--border);
  border-radius: 14px;
  padding: 2rem 1.5rem;
  text-align: center;
  transition: border-color .2s, transform .2s, box-shadow .2s;
  box-shadow: var(--shadow);
}
.feat-card:hover {
  border-color: var(--teal);
  transform: translateY(-3px);
  box-shadow: var(--shadow2);
}
.feat-icon  { font-size: 2.2rem; margin-bottom: .8rem; }
.feat-title { font-family: var(--font-h); font-weight: 700; font-size: 1rem; margin-bottom: .4rem; color: var(--ink); }
.feat-desc  { font-size: .83rem; color: var(--ink3); line-height: 1.6; }

/* ══════════════════════════════
   THUMBNAIL
══════════════════════════════ */
.thumb-wrap {
  border-radius: var(--r);
  overflow: hidden;
  border: 1.5px solid var(--border);
  background: var(--cream2);
  box-shadow: var(--shadow);
}
.thumb-wrap img { width: 100%; display: block; }
.thumb-meta { padding: .8rem 1rem; background: var(--cream2); }
.vid-badge {
  font-family: var(--font-m);
  font-size: .72rem;
  background: #fff;
  border: 1px solid var(--border);
  color: var(--ink3);
  padding: .22rem .65rem;
  border-radius: 6px;
}

/* ══════════════════════════════
   PILLS / BADGES
══════════════════════════════ */
.pill {
  display: inline-flex;
  align-items: center;
  gap: .35rem;
  font-family: var(--font-b);
  font-size: .76rem;
  font-weight: 500;
  padding: .3rem .8rem;
  border-radius: 100px;
}
.pill-teal   { background: var(--teal-lt);  border: 1px solid rgba(26,107,107,.25); color: var(--teal); }
.pill-amber  { background: var(--amber-lt); border: 1px solid rgba(217,119, 6,.25); color: var(--amber); }
.pill-rose   { background: var(--rose-lt);  border: 1px solid rgba(190, 18,60,.2);  color: var(--rose); }
.pill-ink    { background: var(--cream2);   border: 1px solid var(--border);        color: var(--ink2); }

/* ══════════════════════════════
   PIPELINE STEPS (sidebar)
══════════════════════════════ */
.step-item {
  display: flex;
  align-items: center;
  gap: .6rem;
  padding: .4rem .5rem;
  font-size: .81rem;
  color: #78716c;
  border-radius: 7px;
  font-family: var(--font-b);
}
.step-item.done    { color: #4ade80; }
.step-item.running { color: #5eead4; background: rgba(94,234,212,.08); }
.step-icon { width: 1rem; text-align: center; font-size: .78rem; }

/* ══════════════════════════════
   CHAT BUBBLES
══════════════════════════════ */
.chat-row-user {
  display: flex; justify-content: flex-end;
  margin: .6rem 0; animation: fadeUp .22s ease;
}
.chat-row-bot {
  display: flex; justify-content: flex-start;
  margin: .6rem 0; animation: fadeUp .22s ease;
}
.bubble-user {
  background: var(--teal);
  color: #fff;
  padding: .75rem 1.15rem;
  border-radius: 20px 20px 4px 20px;
  max-width: 74%;
  font-size: .9rem;
  line-height: 1.6;
  box-shadow: 0 2px 10px rgba(26,107,107,.28);
}
.bubble-bot {
  background: #fff;
  border: 1.5px solid var(--border);
  color: var(--ink2);
  padding: .75rem 1.15rem;
  border-radius: 20px 20px 20px 4px;
  max-width: 82%;
  font-size: .9rem;
  line-height: 1.75;
  white-space: pre-wrap;
  box-shadow: var(--shadow);
}
@keyframes fadeUp {
  from { opacity: 0; transform: translateY(7px); }
  to   { opacity: 1; transform: translateY(0); }
}

/* ══════════════════════════════
   OUTPUT BOXES
══════════════════════════════ */
.out-box {
  background: #fff;
  border: 1.5px solid var(--border);
  border-radius: var(--r);
  padding: 1.4rem 1.6rem;
  font-size: .9rem;
  line-height: 1.85;
  color: var(--ink2);
  white-space: pre-wrap;
  box-shadow: var(--shadow);
}
.out-box.teal  { border-left: 4px solid var(--teal); }
.out-box.amber { border-left: 4px solid var(--amber); }
.out-box.rose  { border-left: 4px solid var(--rose); }

/* ══════════════════════════════
   QUIZ  — fully hardcoded colours
   (no CSS vars, guaranteed render)
══════════════════════════════ */
.quiz-q-box {
  background: #1c1917;
  border-radius: 10px;
  padding: 1rem 1.25rem;
  margin-bottom: .9rem;
  font-family: 'DM Sans', sans-serif;
  font-size: .97rem;
  font-weight: 600;
  color: #faf7f2;
  line-height: 1.5;
  border-left: 4px solid #2a8f8f;
}

.q-opt {
  display: flex;
  align-items: center;
  gap: .75rem;
  padding: .65rem 1.1rem;
  margin-bottom: .4rem;
  border-radius: 9px;
  font-family: 'DM Sans', sans-serif;
  font-size: .88rem;
  line-height: 1.45;
  border: 1.5px solid #d6cfc4;
  background: #ffffff;
  color: #44403c;
}
.q-opt .opt-lbl {
  min-width: 1.7rem;
  height: 1.7rem;
  border-radius: 50%;
  display: flex;
  align-items: center;
  justify-content: center;
  font-size: .74rem;
  font-weight: 700;
  flex-shrink: 0;
  background: #ede5d8;
  color: #78716c;
}

/* correct */
.q-opt.correct {
  border-color: #059669 !important;
  background: #ecfdf5 !important;
  color: #065f46 !important;
}
.q-opt.correct .opt-lbl {
  background: #059669 !important;
  color: #ffffff !important;
}

/* wrong */
.q-opt.wrong {
  border-color: #be123c !important;
  background: #fff1f2 !important;
  color: #9f1239 !important;
}
.q-opt.wrong .opt-lbl {
  background: #be123c !important;
  color: #ffffff !important;
}

/* neutral (other options after submit) */
.q-opt.neutral {
  opacity: .4;
}

/* feedback line */
.quiz-fb {
  display: inline-flex;
  align-items: center;
  gap: .4rem;
  font-family: 'DM Sans', sans-serif;
  font-size: .82rem;
  font-weight: 600;
  padding: .28rem .8rem;
  border-radius: 100px;
  margin-top: .3rem;
  margin-bottom: .2rem;
}
.quiz-fb.ok  { background: #d1fae5; color: #065f46; border: 1px solid #6ee7b7; }
.quiz-fb.bad { background: #ffe4e6; color: #9f1239; border: 1px solid #fda4af; }

.quiz-score-box {
  text-align: center;
  padding: 1.6rem 2rem;
  background: #fff;
  border: 1.5px solid #d6cfc4;
  border-radius: 14px;
  margin-top: 1.2rem;
  font-family: 'DM Sans', sans-serif;
  font-size: 1.05rem;
  font-weight: 600;
  color: #1c1917;
  box-shadow: 0 2px 16px rgba(28,25,23,.08);
}

/* ══════════════════════════════
   METRICS
══════════════════════════════ */
[data-testid="stMetric"] {
  background: #fff;
  border: 1.5px solid var(--border);
  border-radius: 10px;
  padding: .9rem 1.1rem;
  box-shadow: var(--shadow);
}
[data-testid="stMetricLabel"] { color: var(--ink3) !important; font-size: .78rem !important; }
[data-testid="stMetricValue"] { color: var(--ink) !important; font-size: 1.5rem !important; font-weight: 700 !important; }

/* ══════════════════════════════
   DOWNLOAD BUTTONS
══════════════════════════════ */
[data-testid="stDownloadButton"] button {
  background: #fff !important;
  color: var(--teal) !important;
  border: 1.5px solid var(--border) !important;
  border-radius: 8px !important;
  font-family: var(--font-b) !important;
  font-size: .85rem !important;
  font-weight: 500 !important;
  transition: all .18s !important;
  box-shadow: var(--shadow);
}
[data-testid="stDownloadButton"] button:hover {
  background: var(--teal-lt) !important;
  border-color: var(--teal) !important;
  transform: translateY(-1px) !important;
}

/* ══════════════════════════════
   MISC
══════════════════════════════ */
hr { border-color: var(--cream3) !important; margin: 1.3rem 0 !important; }
.stSpinner > div { border-top-color: var(--teal) !important; }
textarea {
  background: #fff !important;
  border: 1.5px solid var(--border) !important;
  border-radius: 10px !important;
  color: var(--ink2) !important;
  font-family: var(--font-m) !important;
  font-size: .78rem !important;
  line-height: 1.6 !important;
}
.sec-label {
  font-family: var(--font-b);
  font-size: .72rem;
  font-weight: 600;
  letter-spacing: .13em;
  text-transform: uppercase;
  color: var(--ink3);
  margin-bottom: .75rem;
}
.sb-logo {
  font-family: 'Playfair Display', serif;
  font-size: 1.3rem;
  font-weight: 800;
  color: #faf7f2 !important;
  letter-spacing: -.01em;
  margin-bottom: .2rem;
}
.sb-sub {
  font-size: .78rem;
  color: #a8a29e !important;
  line-height: 1.5;
  margin-bottom: 1.5rem;
}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════
STEPS = ["Fetch Transcript","Build Document","Chunk & Index",
         "Create Embeddings","Build Vector DB","Initialize Chain"]

_DEFAULTS: dict = {
    "step_status":     {s:"pending" for s in STEPS},
    "chain":           None,
    "transcript_text": None,
    "video_id":        None,
    "chat_history":    [],
    "quiz_questions":  [],
    "quiz_answers":    {},
    "quiz_submitted":  {},
    "quiz_generated":  False,
    "summary_text":    None,
    "keypoints_text":  None,
    "active_url":      "",
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ══════════════════════════════════════════════════════════
# EXPORT HELPERS
# ══════════════════════════════════════════════════════════
def transcript_to_pdf(text: str, video_id: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=1.1*inch, rightMargin=1.1*inch,
        topMargin=1*inch, bottomMargin=1*inch)

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("T", parent=styles["Title"],
        fontSize=22, textColor=colors.HexColor("#1a6b6b"),
        spaceAfter=4, alignment=TA_CENTER,
        fontName="Helvetica-Bold")
    meta_s = ParagraphStyle("M", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#78716c"),
        spaceAfter=20, alignment=TA_CENTER)
    body_s = ParagraphStyle("B", parent=styles["Normal"],
        fontSize=10.5, leading=17,
        textColor=colors.HexColor("#1c1917"),
        spaceAfter=6)

    story = [
        Paragraph("TubeIQ — Video Transcript", title_s),
        Paragraph(f"Video ID: {video_id}", meta_s),
    ]
    for chunk in [text[i:i+700] for i in range(0, len(text), 700)]:
        safe = chunk.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        story.append(Paragraph(safe, body_s))
        story.append(Spacer(1, 3))

    doc.build(story)
    return buf.getvalue()


def transcript_to_docx(text: str, video_id: str) -> bytes:
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDoc()
    # Margins
    for section in doc.sections:
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    title = doc.add_heading("TubeIQ — Video Transcript", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.runs[0]
    tr.font.color.rgb = RGBColor(0x1a, 0x6b, 0x6b)
    tr.font.size = Pt(22)

    meta = doc.add_paragraph(f"Video ID: {video_id}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(0x78, 0x71, 0x6c)
    meta.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    pos = 0
    while pos < len(text):
        end = min(pos + 600, len(text))
        if end < len(text):
            sp = text.rfind(" ", pos, end)
            if sp > pos: end = sp
        p = doc.add_paragraph(text[pos:end].strip())
        p.runs[0].font.size = Pt(10.5)
        pos = end + 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════
# CORE HELPERS
# ══════════════════════════════════════════════════════════
def extract_video_id(url: str) -> str:
    url = url.strip()
    if "v=" in url:
        return url.split("v=")[1].split("&")[0]
    if "youtu.be/" in url:
        return url.split("youtu.be/")[1].split("?")[0]
    return url

def reset_state():
    for k, v in _DEFAULTS.items():
        st.session_state[k] = v

def get_transcript(video_id: str):
    try:
        from youtube_transcript_api.proxies import GenericProxyConfig
        key = os.environ.get("SCRAPER_API_KEY", "")
        proxy_url = f"http://scraperapi:{key}@proxy-server.scraperapi.com:8001"
        cfg = GenericProxyConfig(http_url=proxy_url, https_url=proxy_url)
        sess = requests.Session()
        sess.verify = False
        requests.packages.urllib3.disable_warnings()
        api = YouTubeTranscriptApi(proxy_config=cfg, http_client=sess)
        try:    fetched = api.fetch(video_id, languages=["en"])
        except: fetched = api.fetch(video_id)
        return fetched.to_raw_data()
    except TranscriptsDisabled:
        st.error("No captions available for this video.")
        return None
    except Exception as e:
        st.error(f"Transcript error: {e}")
        return None

def build_document(tlist):
    full, omap, cursor = "", [], 0
    for c in tlist:
        t = c["text"].strip() + " "
        omap.append({"char_start":cursor,"char_end":cursor+len(t),
                     "start_time":c["start"],"end_time":c["start"]+c["duration"]})
        full += t; cursor += len(t)
    return Document(page_content=full, metadata={"offset_map":omap}), full

def split_with_timestamps(doc):
    total = len(doc.page_content)
    cs = 1500 if total > 100_000 else 1000
    co = 300  if total > 100_000 else 200
    splitter = RecursiveCharacterTextSplitter(chunk_size=cs, chunk_overlap=co, add_start_index=True)
    chunks = splitter.split_documents([doc])
    omap = doc.metadata["offset_map"]
    out = []
    for ch in chunks:
        s = ch.metadata["start_index"]; e = s + len(ch.page_content)
        st_t = en_t = None
        for o in omap:
            if o["char_end"] >= s and st_t is None: st_t = o["start_time"]
            if o["char_start"] <= e: en_t = o["end_time"]
        ch.metadata = {"start":st_t,"end":en_t}
        out.append(ch)
    return out

@st.cache_resource(show_spinner=False)
def build_vector_store(_key, docs):
    emb = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return FAISS.from_documents(docs, emb)

def format_docs(docs):
    return "\n\n---\n\n".join(
        f"[Chunk {i}] Start={d.metadata.get('start','?')}s End={d.metadata.get('end','?')}s\n{d.page_content}"
        for i, d in enumerate(docs, 1)
    )

def build_chain(retriever):
    llm = ChatGroq(model="moonshotai/kimi-k2-instruct-0905", temperature=0.2, streaming=True)
    prompt = PromptTemplate(
        template="""You are TubeIQ. Answer ONLY from the transcript context below.
- Always cite timestamp ranges.
- If not in context say exactly: "I don't know based on the provided transcript."

Context:
{context}

Question: {question}

Answer: <explanation>
Sources:
- Timestamp: <start - end>
""",
        input_variables=["context","question"],
    )
    return (
        RunnableParallel({"context": retriever | RunnableLambda(format_docs), "question": RunnablePassthrough()})
        | prompt | llm | StrOutputParser()
    )

def extract_earliest_ts(text: str) -> int:
    m = re.findall(r"Timestamp:\s*([\d\.]+)\s*-\s*([\d\.]+)", text)
    return int(min(float(s) for s, _ in m)) if m else 0

def stream_to_box(prompt_text: str, style: str = "teal") -> str:
    box = st.empty(); full = ""
    for chunk in st.session_state.chain.stream(prompt_text):
        full += chunk
        box.markdown(f'<div class="out-box {style}">{full}▌</div>', unsafe_allow_html=True)
    box.markdown(f'<div class="out-box {style}">{full}</div>', unsafe_allow_html=True)
    return full

def parse_quiz(raw: str) -> list:
    questions = []
    blocks = re.split(r"\n?Q\d+[\.\)]\s*", raw.strip())
    for block in blocks:
        if not block.strip(): continue
        lines = [l.strip() for l in block.strip().splitlines() if l.strip()]
        if not lines: continue
        q_text = lines[0]; options = {}; correct = None
        for line in lines[1:]:
            m = re.match(r"^([A-D])[\)\.]\s*(.+)", line)
            if m: options[m.group(1)] = m.group(2)
            am = re.match(r"(?:Answer|Correct)[:\s]+([A-D])", line, re.IGNORECASE)
            if am: correct = am.group(1).upper()
        if q_text and len(options) >= 2 and correct:
            questions.append({"q":q_text,"options":options,"answer":correct})
    return questions


# ══════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown('<div class="sb-logo">TubeIQ</div>', unsafe_allow_html=True)
    st.markdown('<div class="sb-sub">AI video intelligence — chat, summarize and quiz any YouTube video.</div>', unsafe_allow_html=True)

    # URL input with explicit Load button (Enter also works via form)
    with st.form("url_form", clear_on_submit=False):
        video_url_input = st.text_input(
            "YouTube URL",
            value=st.session_state.active_url,
            placeholder="https://youtube.com/watch?v=...",
        )
        url_submitted = st.form_submit_button("▶  Analyse Video", use_container_width=True)

    if url_submitted and video_url_input.strip():
        new_id = extract_video_id(video_url_input.strip())
        if new_id != st.session_state.video_id:
            reset_state()
            st.session_state.active_url = video_url_input.strip()
            st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    if st.session_state.chain:
        st.markdown('<span class="pill pill-teal">● Ready</span>', unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("↺  Load different video", key="reset_btn"):
            reset_state()
            st.rerun()

    st.markdown("<hr style='border-color:#292524;margin:1.2rem 0'>", unsafe_allow_html=True)
    st.markdown('<div style="font-size:.72rem;font-weight:600;letter-spacing:.12em;text-transform:uppercase;color:#57534e;margin-bottom:.5rem">Pipeline</div>', unsafe_allow_html=True)

    icons = {"pending":"○","running":"◌","done":"●"}
    for step in STEPS:
        status = st.session_state.step_status[step]
        st.markdown(
            f'<div class="step-item {status}"><span class="step-icon">{icons[status]}</span>{step}</div>',
            unsafe_allow_html=True,
        )


# ══════════════════════════════════════════════════════════
# HERO
# ══════════════════════════════════════════════════════════
st.markdown("""
<div class="hero-wrap">
  <div class="hero-eyebrow">Groq · LangChain · FAISS · Streamlit</div>
  <div class="hero-title">Tube<span>IQ</span></div>
  <div class="hero-sub">Ask anything about any YouTube video. Instant answers, structured summaries, and interactive quizzes — all from the transcript.</div>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# LANDING
# ══════════════════════════════════════════════════════════
if not st.session_state.active_url:
    c1, c2, c3, c4 = st.columns(4)
    for col, icon, title, desc in [
        (c1,"💬","Smart Chat",      "Ask follow-up questions and get cited, timestamped answers."),
        (c2,"📋","Auto Summary",    "Structured summaries with key points and notable timestamps."),
        (c3,"🧠","Interactive Quiz","Take a generated quiz and get instant per-answer feedback."),
        (c4,"📄","Transcript","View and download the full transcript as TXT, PDF or Word."),
    ]:
        with col:
            st.markdown(f"""<div class="feat-card">
                <div class="feat-icon">{icon}</div>
                <div class="feat-title">{title}</div>
                <div class="feat-desc">{desc}</div>
            </div>""", unsafe_allow_html=True)
    st.stop()


# ══════════════════════════════════════════════════════════
# VIDEO ID
# ══════════════════════════════════════════════════════════
video_id = extract_video_id(st.session_state.active_url)


# ══════════════════════════════════════════════════════════
# TOP ROW — thumbnail + pipeline status + quick actions
# ══════════════════════════════════════════════════════════
col_thumb, col_main = st.columns([5, 7], gap="large")

with col_thumb:
    st.markdown(f"""<div class="thumb-wrap">
        <img src="https://img.youtube.com/vi/{video_id}/maxresdefault.jpg"
             onerror="this.src='https://img.youtube.com/vi/{video_id}/hqdefault.jpg'"/>
        <div class="thumb-meta"><span class="vid-badge">{video_id}</span></div>
    </div>""", unsafe_allow_html=True)

with col_main:
    progress_slot = st.empty()

    if st.session_state.chain is None:
        def mark(step, status): st.session_state.step_status[step] = status
        def info(msg): progress_slot.markdown(f'<span class="pill pill-amber">⟳ {msg}</span>', unsafe_allow_html=True)

        mark("Fetch Transcript","running"); info("Fetching transcript…")
        tlist = get_transcript(video_id)
        if not tlist: st.stop()
        mark("Fetch Transcript","done")

        mark("Build Document","running"); info("Building document…")
        doc, full_text = build_document(tlist)
        st.session_state.transcript_text = full_text
        mark("Build Document","done")

        mark("Chunk & Index","running"); info("Chunking…")
        final_docs = split_with_timestamps(doc)
        mark("Chunk & Index","done")

        mark("Create Embeddings","running"); info("Creating embeddings…")
        mark("Create Embeddings","done")
        mark("Build Vector DB","running"); info("Building vector DB…")
        vs = build_vector_store(video_id, final_docs)
        retriever = vs.as_retriever(search_type="similarity", search_kwargs={"k":5})
        mark("Build Vector DB","done")

        mark("Initialize Chain","running"); info("Initializing chain…")
        st.session_state.chain    = build_chain(retriever)
        st.session_state.video_id = video_id
        mark("Initialize Chain","done")
        progress_slot.markdown('<span class="pill pill-teal">✓ Ready — choose a tab below</span>', unsafe_allow_html=True)
    else:
        progress_slot.markdown('<span class="pill pill-teal">✓ Video loaded and ready</span>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Quick-action row
    qa1, qa2, qa3 = st.columns(3)
    with qa1: do_summary   = st.button("📋 Summarize video",  key="qa_s")
    with qa2: do_quiz      = st.button("🧠 Generate quiz",     key="qa_q")
    with qa3: do_keypoints = st.button("🔑 Extract key points",key="qa_k")


# ══════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════
tab_chat, tab_tools, tab_quiz, tab_transcript = st.tabs(
    ["💬  Chat", "🛠  Tools", "🧠  Quiz", "📄  Transcript"]
)


# ──────────────────────────────────────────────────────────
# TAB 1 — CHAT
# Form prevents the question-loop bug on rerun
# ──────────────────────────────────────────────────────────
with tab_chat:
    # Render history
    for msg in st.session_state.chat_history:
        css    = "chat-row-user" if msg["role"] == "user" else "chat-row-bot"
        bubble = "bubble-user"   if msg["role"] == "user" else "bubble-bot"
        st.markdown(
            f'<div class="{css}"><div class="{bubble}">{msg["content"]}</div></div>',
            unsafe_allow_html=True,
        )

    # Input form — clear_on_submit prevents rerun loop
    with st.form("chat_form", clear_on_submit=True):
        fc1, fc2 = st.columns([8, 1])
        with fc1:
            question = st.text_input(
                "q", placeholder="Ask anything about this video…",
                label_visibility="collapsed"
            )
        with fc2:
            chat_submit = st.form_submit_button("Send")

    if chat_submit and question.strip():
        q = question.strip()
        st.session_state.chat_history.append({"role":"user","content":q})
        st.markdown(f'<div class="chat-row-user"><div class="bubble-user">{q}</div></div>', unsafe_allow_html=True)

        with st.spinner(""):
            box, full = st.empty(), ""
            for chunk in st.session_state.chain.stream(q):
                full += chunk
                box.markdown(
                    f'<div class="chat-row-bot"><div class="bubble-bot">{full}▌</div></div>',
                    unsafe_allow_html=True,
                )
            box.markdown(
                f'<div class="chat-row-bot"><div class="bubble-bot">{full}</div></div>',
                unsafe_allow_html=True,
            )

        st.session_state.chat_history.append({"role":"assistant","content":full})

        ts = extract_earliest_ts(full)
        if ts > 0:
            st.markdown('<div class="sec-label" style="margin-top:1.2rem">🎯 Relevant segment</div>', unsafe_allow_html=True)
            st.components.v1.iframe(
                f"https://www.youtube.com/embed/{st.session_state.video_id}?start={ts}",
                height=300,
            )

    if st.session_state.chat_history:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🗑  Clear conversation", key="clr"):
            st.session_state.chat_history = []
            st.rerun()


# ──────────────────────────────────────────────────────────
# TAB 2 — TOOLS
# ──────────────────────────────────────────────────────────
with tab_tools:
    t1, t2 = st.columns(2)
    with t1: trig_sum = st.button("📋 Summarize",   key="t_s")
    with t2: trig_kp  = st.button("🔑 Key Points",  key="t_k")

    if do_summary or trig_sum:
        st.markdown('<div class="sec-label">📋 Summary</div>', unsafe_allow_html=True)
        with st.spinner("Generating…"):
            r = stream_to_box(
                "Provide a well-structured summary covering: "
                "1) Main topic and purpose, 2) Key arguments, "
                "3) Important takeaways, 4) Notable timestamps to revisit.",
                "teal"
            )
            st.session_state.summary_text = r
    elif st.session_state.summary_text:
        st.markdown('<div class="sec-label">📋 Summary</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="out-box teal">{st.session_state.summary_text}</div>', unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    if do_keypoints or trig_kp:
        st.markdown('<div class="sec-label">🔑 Key Points</div>', unsafe_allow_html=True)
        with st.spinner("Extracting…"):
            r = stream_to_box(
                "List the 8 most important key points. "
                "One concise bullet per point with the relevant timestamp range.",
                "amber"
            )
            st.session_state.keypoints_text = r
    elif st.session_state.keypoints_text:
        st.markdown('<div class="sec-label">🔑 Key Points</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="out-box amber">{st.session_state.keypoints_text}</div>', unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown('<div class="sec-label">🔍 Keyword search</div>', unsafe_allow_html=True)
    kw = st.text_input("Keyword", placeholder="e.g. climate change", key="kw")
    if kw and st.session_state.transcript_text:
        count = st.session_state.transcript_text.lower().count(kw.lower())
        tag = "pill-teal" if count else "pill-rose"
        msg = f'✓ "{kw}" found {count} time(s)' if count else f'✗ "{kw}" not found'
        st.markdown(f'<span class="pill {tag}">{msg}</span>', unsafe_allow_html=True)


# ──────────────────────────────────────────────────────────
# TAB 3 — INTERACTIVE QUIZ
# All HTML uses hardcoded hex colours — no CSS vars
# ──────────────────────────────────────────────────────────
with tab_quiz:
    gen_col, _ = st.columns([2, 5])
    with gen_col:
        trig_gen = st.button("🧠 Generate quiz", key="q_gen")

    if do_quiz or trig_gen:
        st.session_state.quiz_questions = []
        st.session_state.quiz_answers   = {}
        st.session_state.quiz_submitted = {}
        st.session_state.quiz_generated = False

        with st.spinner("Generating questions…"):
            raw, box = "", st.empty()
            for chunk in st.session_state.chain.stream(
                "Generate exactly 5 multiple choice questions from the video transcript.\n"
                "Strict format for every question:\n\n"
                "Q1. <question text>\n"
                "A) <option>\nB) <option>\nC) <option>\nD) <option>\n"
                "Answer: <letter>\n\n"
                "Q2. ... and so on. Vary difficulty."
            ):
                raw += chunk
                box.markdown(f'<div class="out-box teal">{raw}▌</div>', unsafe_allow_html=True)
            box.empty()

        parsed = parse_quiz(raw)
        if parsed:
            st.session_state.quiz_questions = parsed
            st.session_state.quiz_generated = True
        else:
            st.warning("Could not parse quiz. Try generating again.")

    # ── Render quiz ──
    if st.session_state.quiz_generated and st.session_state.quiz_questions:
        questions     = st.session_state.quiz_questions
        total         = len(questions)
        correct_count = 0

        for i, qd in enumerate(questions):
            q_text  = qd["q"]
            options = qd["options"]
            answer  = qd["answer"]
            chosen  = st.session_state.quiz_answers.get(i)
            done    = st.session_state.quiz_submitted.get(i, False)

            # Question box — dark background, light text, hardcoded
            st.markdown(
                f'<div class="quiz-q-box">Q{i+1}. {q_text}</div>',
                unsafe_allow_html=True,
            )

            if not done:
                # Clickable answer buttons
                for letter, text in options.items():
                    if st.button(f"{letter})  {text}", key=f"q{i}_{letter}"):
                        st.session_state.quiz_answers[i]   = letter
                        st.session_state.quiz_submitted[i] = True
                        st.rerun()
            else:
                # Styled result rows — all colours hardcoded
                for letter, text in options.items():
                    if letter == answer:
                        css = "correct"; icon = "✓"
                    elif letter == chosen:
                        css = "wrong";   icon = "✗"
                    else:
                        css = "neutral"; icon = ""
                    st.markdown(
                        f'<div class="q-opt {css}">'
                        f'<span class="opt-lbl">{letter}</span>'
                        f'<span>{icon} {text}</span>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

                if chosen == answer:
                    correct_count += 1
                    st.markdown('<span class="quiz-fb ok">✓ Correct!</span>', unsafe_allow_html=True)
                else:
                    st.markdown(
                        f'<span class="quiz-fb bad">✗ Correct: {answer}) {options.get(answer,"")}</span>',
                        unsafe_allow_html=True,
                    )

            st.markdown("<hr>", unsafe_allow_html=True)

        # Score
        answered = len(st.session_state.quiz_submitted)
        if answered == total:
            pct   = int((correct_count / total) * 100)
            grade = "🏆 Excellent!" if pct >= 80 else ("👍 Good job!" if pct >= 60 else "📖 Keep reviewing!")
            st.markdown(f"""<div class="quiz-score-box">
                {grade} &nbsp;&nbsp; Score: <strong>{correct_count}/{total}</strong> ({pct}%)
            </div>""", unsafe_allow_html=True)
        else:
            st.markdown(f'<span class="pill pill-ink">{answered}/{total} answered</span>', unsafe_allow_html=True)

    elif not st.session_state.quiz_generated:
        st.markdown("""<div class="card" style="text-align:center;padding:3rem 2rem;box-shadow:none">
            <div style="font-size:2.6rem;margin-bottom:.9rem">🧠</div>
            <div style="font-family:'Playfair Display',serif;font-weight:700;font-size:1.05rem;margin-bottom:.5rem;color:#1c1917">Interactive Quiz</div>
            <div style="color:#78716c;font-size:.85rem;line-height:1.65">
                Click <strong>Generate quiz</strong> to create 5 questions from this video.<br>
                Click an option to lock in your answer and get instant feedback.
            </div>
        </div>""", unsafe_allow_html=True)


# ──────────────────────────────────────────────────────────
# TAB 4 — TRANSCRIPT + DOWNLOADS
# ──────────────────────────────────────────────────────────
with tab_transcript:
    if st.session_state.transcript_text:
        txt = st.session_state.transcript_text
        vid = st.session_state.video_id or "video"

        m1, m2, m3 = st.columns(3)
        m1.metric("Words",        f"{len(txt.split()):,}")
        m2.metric("Characters",   f"{len(txt):,}")
        m3.metric("Est. length",  f"~{round(len(txt)/14/60)} min")

        st.markdown("<br>", unsafe_allow_html=True)
        st.text_area("Full transcript", txt, height=400, label_visibility="visible")

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="sec-label">⬇ Download transcript</div>', unsafe_allow_html=True)

        dl1, dl2, dl3 = st.columns(3)

        with dl1:
            st.download_button(
                label="📄 Plain text (.txt)",
                data=txt,
                file_name=f"transcript_{vid}.txt",
                mime="text/plain",
                use_container_width=True,
            )

        with dl2:
            try:
                pdf_bytes = transcript_to_pdf(txt, vid)
                st.download_button(
                    label="📕 PDF document (.pdf)",
                    data=pdf_bytes,
                    file_name=f"transcript_{vid}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.caption(f"PDF unavailable — install reportlab: {e}")

        with dl3:
            try:
                docx_bytes = transcript_to_docx(txt, vid)
                st.download_button(
                    label="📘 Word document (.docx)",
                    data=docx_bytes,
                    file_name=f"transcript_{vid}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.caption(f"DOCX unavailable — install python-docx: {e}")
    else:
        st.info("Load a video first to view and download its transcript.")